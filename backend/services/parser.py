import logging
import fitz  # PyMuPDF
from docx import Document
import io
import pytesseract
from PIL import Image

logger = logging.getLogger("vibeonjob.services.parser")

from services.ocr_engine import extract_text_from_image

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file.

    Entry point that delegates to the best available extraction strategy.
    Currently uses PyMuPDF native text extraction, which is optimal for
    text-selectable PDFs.
    """
    return parse_pdf_with_pymupdf(file_bytes)


def parse_pdf_with_pymupdf(file_bytes: bytes) -> str:
    """Extract text from a text-selectable PDF using PyMuPDF's native API.

    Optimised for accuracy on digitally-created (non-scanned) PDFs:
    - Uses ``get_text("text")`` which preserves reading order via PyMuPDF's
      internal layout analysis.
    - Sorts extracted blocks by vertical then horizontal position to maintain
      correct reading flow across multi-column layouts.
    - Strips trailing whitespace per page and skips blank pages.
    """
    logger.debug("Initializing PyMuPDF native text extraction")
    out_parts: list[str] = []
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            logger.debug(f"PDF opened successfully. Number of pages: {len(doc)}")
            for i, page in enumerate(doc):
                # Extract text blocks: each block is (x0, y0, x1, y1, "text", block_no, block_type)
                blocks = page.get_text("blocks")
                # Keep only text blocks (block_type == 0) and sort by position (top-to-bottom, left-to-right)
                text_blocks = sorted(
                    [b for b in blocks if b[6] == 0],
                    key=lambda b: (b[1], b[0]),
                )
                page_text = "\n".join(b[4].strip() for b in text_blocks if b[4].strip())

                if page_text:
                    out_parts.append(page_text)
                else:
                    logger.debug(f"Page {i + 1} yielded no selectable text")

        result = "\n\n".join(out_parts)
        logger.debug(f"Extracted {len(result)} characters from PDF")
        return result
    except Exception as e:
        logger.error(f"Failed to parse PDF document: {e}")
        raise


def parse_pdf_with_ocr(file_bytes: bytes) -> str:
    """Extract text from a scanned/image-based PDF using OCR.

    Renders each page at 300 DPI and runs through the OCR engine pipeline
    (deskew → preprocess → Tesseract).  Falls back to selectable text when
    OCR yields nothing for a given page.

    Note: This path is intended for non-selectable (scanned) PDFs.  For
    text-selectable documents, ``parse_pdf_with_pymupdf`` is preferred.
    """
    logger.debug("Initializing PyMuPDF parser with OCR pipeline")
    out_parts: list[str] = []
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            logger.debug(f"PDF opened successfully. Number of pages: {len(doc)}")
            for i, page in enumerate(doc):
                selectable = page.get_text() or ""

                try:
                    pix = page.get_pixmap(dpi=300)
                    img_bytes = pix.tobytes("png")
                    pil_img = Image.open(io.BytesIO(img_bytes))
                    ocr_text = extract_text_from_image(pil_img)
                except Exception as e:
                    logger.warning(f"OCR failed on page {i + 1}: {e}")
                    ocr_text = ""

                # Prefer OCR when it captures significantly more content
                if ocr_text.strip() and len(ocr_text) > len(selectable) * 1.5:
                    out_parts.append(ocr_text)
                elif ocr_text.strip() and selectable.strip():
                    out_parts.append(selectable + "\n" + ocr_text)
                elif selectable.strip():
                    out_parts.append(selectable)
                elif ocr_text.strip():
                    out_parts.append(ocr_text)

        return "\n".join(p for p in out_parts if p)
    except Exception as e:
        logger.error(f"Failed to parse PDF via OCR: {e}")
        raise

def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    logger.debug("Initializing python-docx parser")
    try:
        doc = Document(io.BytesIO(file_bytes))
        logger.debug(f"DOCX opened successfully. Number of paragraphs: {len(doc.paragraphs)}")
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Failed to parse DOCX document: {e}")
        raise

def parse_image(file_bytes: bytes) -> str:
    """Extract text from an image file (JPG, PNG)."""
    logger.debug("Initializing Image parser with OCR pipeline")
    try:
        pil_img = Image.open(io.BytesIO(file_bytes))
        text = extract_text_from_image(pil_img)
        return text
    except Exception as e:
        logger.error(f"Failed to parse image document: {e}")
        raise
