import logging
import fitz  # PyMuPDF
from docx import Document
import io
import pytesseract
from PIL import Image

logger = logging.getLogger("vibeonjob.services.parser")

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    logger.debug("Initializing PyMuPDF parser")
    text = ""
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            logger.debug(f"PDF opened successfully. Number of pages: {len(doc)}")
            for i, page in enumerate(doc):
                page_text = page.get_text()

                  # 2. Exhaustive Full-Page OCR Verification
                logger.info(f"Generating full-page 300 DPI visual render for OCR sweep on page {i+1}.")
                try:
                    # Generate a high-resolution visual snapshot of the whole page
                    pix = page.get_pixmap(dpi=300)
                    img_bytes = pix.tobytes("png")
                    pil_img = Image.open(io.BytesIO(img_bytes))
                    
                    # Extract text via Tesseract OCR
                    ocr_text = pytesseract.image_to_string(pil_img)
                    if ocr_text.strip():
                        logger.debug(f"Full-page OCR sweep extracted {len(ocr_text.strip())} chars on page {i+1}")
                        page_text += "\n" + ocr_text
                except Exception as ocr_err:
                    logger.warning(f"Full-page OCR sweep failed on page {i+1} (Tesseract may be missing): {ocr_err}")
                    
                text += page_text + "\n"
                logger.debug(f"Extracted {len(page_text)} chars from page {i+1}")
        return text
    except Exception as e:
        logger.error(f"Failed to parse PDF document: {e}")
        raise

def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    logger.debug("Initializing python-docx parser")
    try:
        doc = Document(io.BytesIO(file_bytes))
        logger.debug(f"DOCX opened successfully. Number of paragraphs: {len(doc.paragraphs)}")
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Failed to parse DOCX document: {e}")
        raise
